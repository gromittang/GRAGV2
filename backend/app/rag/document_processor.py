"""
文档处理
支持 PDF、Word、TXT 等格式
支持图片提取和保存，保持原始顺序
"""
from llama_index.core import Document
from llama_index.core.node_parser import SentenceSplitter
from typing import List, Dict, Any, Optional
import os
import uuid
import re

from app.core.settings import get_industry_settings
from app.config import get_settings

_settings = get_settings()

# 图片存储目录
IMAGES_DIR = os.path.join(_settings.data_dir, "images")


def ensure_images_dir():
    """确保图片目录存在"""
    os.makedirs(IMAGES_DIR, exist_ok=True)


class DocumentProcessor:
    """文档处理器"""

    def __init__(self, industry_type: str = None):
        industry = get_industry_settings(industry_type or _settings.industry_type)
        self.chunk_size = industry.chunk_size
        self.chunk_overlap = industry.chunk_overlap
        self.images = []  # 存储提取的图片信息

    def parse_file(self, file_path: str, metadata: Dict = None) -> List[Document]:
        """
        解析单个文件

        Args:
            file_path: 文件路径
            metadata: 文档元数据

        Returns:
            Document 列表
        """
        ensure_images_dir()
        self.images = []
        docs = []

        file_ext = os.path.splitext(file_path)[1].lower()
        doc_id = metadata.get("document_id", str(uuid.uuid4())) if metadata else str(uuid.uuid4())

        if file_ext == ".txt" or file_ext == ".md":
            # 简单文本文件
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            doc = Document(
                text=content,
                metadata=metadata or {},
                doc_id=doc_id
            )
            docs.append(doc)

        elif file_ext == ".pdf":
            # PDF 使用 PyMuPDF，按页面顺序提取
            docs = self._parse_pdf(file_path, doc_id, metadata)

        elif file_ext == ".docx":
            # Word 文档，按原始顺序提取段落和图片
            docs = self._parse_docx_ordered(file_path, doc_id, metadata)

        else:
            print(f"[DocumentProcessor] 不支持的文件类型: {file_ext}")

        return docs

    def _parse_pdf(self, file_path: str, doc_id: str, metadata: Dict = None) -> List[Document]:
        """解析PDF文件，按页面顺序提取文本和图片"""
        import fitz  # PyMuPDF

        pdf_doc = fitz.open(file_path)
        docs = []
        all_content = []  # 按顺序存储内容
        image_count = 0

        for page_num, page in enumerate(pdf_doc):
            # 获取页面文本块（保持顺序）
            blocks = page.get_text("blocks")

            # 收集页面内所有元素的位置信息
            page_elements = []

            # 添加文本块
            for block in blocks:
                if block[4]:  # 有文本内容
                    page_elements.append({
                        "type": "text",
                        "y": block[1],  # top position
                        "content": block[4]
                    })

            # 添加图片（获取位置）
            images = page.get_images(full=True)
            for img_index, img_info in enumerate(images):
                try:
                    xref = img_info[0]
                    base_image = pdf_doc.extract_image(xref)

                    if base_image:
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]

                        # 保存图片
                        image_filename = f"{doc_id}_p{page_num}_i{img_index}.{image_ext}"
                        image_path = os.path.join(IMAGES_DIR, image_filename)

                        with open(image_path, "wb") as f:
                            f.write(image_bytes)

                        image_url = f"/images/{image_filename}"
                        self.images.append({
                            "filename": image_filename,
                            "path": image_path,
                            "url": image_url,
                            "page": page_num,
                            "index": img_index
                        })

                        # 尝试获取图片位置（从页面对象）
                        img_y = img_info[1] if len(img_info) > 1 else float('inf')
                        page_elements.append({
                            "type": "image",
                            "y": img_y,
                            "content": f"\n[IMG]{image_url}|图片{image_count + 1}[/IMG]\n"
                        })
                        image_count += 1

                except Exception as e:
                    print(f"[PDF] 图片提取失败: {e}")

            # 按y位置排序（模拟原始顺序）
            page_elements.sort(key=lambda x: x["y"] if x["y"] else 0)

            # 添加到总内容
            for elem in page_elements:
                all_content.append(elem["content"])

        pdf_doc.close()

        full_text = "".join(all_content)

        doc = Document(
            text=full_text,
            metadata={**(metadata or {}), "source": file_path, "type": "pdf", "images": len(self.images)},
            doc_id=doc_id
        )
        docs.append(doc)

        print(f"[PDF] 解析完成: {len(full_text)} 字符, {len(self.images)} 张图片")
        return docs

    def _parse_docx_ordered(self, file_path: str, doc_id: str, metadata: Dict = None) -> List[Document]:
        """解析DOCX文件，按原始顺序提取段落和图片"""
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
        from lxml import etree

        docx_doc = DocxDocument(file_path)
        docs = []
        all_content = []
        image_count = 0

        # 构建图片关系映射：rId -> image_data
        image_rels = {}
        for rel in docx_doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_bytes = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    image_ext = content_type.split("/")[-1]
                    if image_ext == "jpeg":
                        image_ext = "jpg"
                    image_rels[rel.rId] = {
                        "bytes": image_bytes,
                        "ext": image_ext
                    }
                except Exception as e:
                    print(f"[DOCX] 图片关系提取失败: {e}")

        # 遍历文档主体元素（保持顺序）
        body = docx_doc.element.body

        for elem in body.iterchildren():
            # 处理段落
            if elem.tag == qn('w:p'):
                para_text = ""
                # 提取段落文本
                for text_elem in elem.findall('.//' + qn('w:t')):
                    if text_elem.text:
                        para_text += text_elem.text

                # 检查段落中是否有嵌入图片
                drawings = elem.findall('.//' + qn('a:blip'))
                for drawing in drawings:
                    # 获取图片引用ID
                    rId = drawing.get(qn('r:embed'))
                    if rId and rId in image_rels:
                        img_data = image_rels[rId]

                        # 保存图片
                        image_filename = f"{doc_id}_i{image_count}.{img_data['ext']}"
                        image_path = os.path.join(IMAGES_DIR, image_filename)

                        with open(image_path, "wb") as f:
                            f.write(img_data["bytes"])

                        image_url = f"/images/{image_filename}"
                        self.images.append({
                            "filename": image_filename,
                            "path": image_path,
                            "url": image_url,
                            "index": image_count
                        })

                        # 在当前位置插入图片标记
                        all_content.append(f"\n[IMG]{image_url}|图片{image_count + 1}[/IMG]\n")
                        image_count += 1

                # 添加段落文本（如果有内容）
                if para_text.strip():
                    all_content.append(para_text + "\n")

            # 处理表格
            elif elem.tag == qn('w:tbl'):
                table_text = ""
                for row in elem.findall('.//' + qn('w:tr')):
                    row_text = []
                    for cell in row.findall('.//' + qn('w:tc')):
                        cell_text = ""
                        for text_elem in cell.findall('.//' + qn('w:t')):
                            if text_elem.text:
                                cell_text += text_elem.text
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text += "| " + " | ".join(row_text) + " |\n"

                if table_text:
                    all_content.append("\n" + table_text + "\n")

        full_text = "".join(all_content)

        doc = Document(
            text=full_text,
            metadata={**(metadata or {}), "source": file_path, "type": "docx", "images": len(self.images)},
            doc_id=doc_id
        )
        docs.append(doc)

        print(f"[DOCX] 解析完成: {len(full_text)} 字符, {len(self.images)} 张图片（按原始顺序）")
        return docs

    def split_documents(self, documents: List[Document]) -> List:
        """
        分块文档（保留图片标记完整性）

        Args:
            documents: 原始文档列表

        Returns:
            分块后的节点列表
        """
        splitter = SentenceSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap
        )

        nodes = splitter.get_nodes_from_documents(documents)

        # 确保图片标记不被截断
        for i, node in enumerate(nodes):
            text = node.text
            # 检查不完整的开始标记
            if "[/IMG]" in text and "[IMG]" not in text:
                # 移除不完整的结束标记
                text = text.replace("[/IMG]", "")
                node.text = text
            # 检查不完整的结束标记
            elif "[IMG]" in text and "[/IMG]" not in text:
                # 找到完整的IMG标签并保留，或者移除不完整的
                start = text.find("[IMG]")
                node.text = text[:start]

        print(f"[DocumentProcessor] 分块完成: {len(documents)} 文档 -> {len(nodes)} 节点")

        return nodes

    def process_file(self, file_path: str, metadata: Dict = None) -> List:
        """
        处理单个文件的完整流程

        Args:
            file_path: 文件路径
            metadata: 元数据

        Returns:
            分块后的节点
        """
        docs = self.parse_file(file_path, metadata)
        nodes = self.split_documents(docs)
        return nodes

    def get_images(self) -> List[Dict]:
        """获取提取的图片列表"""
        return self.images